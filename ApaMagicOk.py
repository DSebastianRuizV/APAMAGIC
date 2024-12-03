import re
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
from docx.oxml.ns import qn

def formatear_y_organizar_referencias(parrafos):
    referencias = []
    for parrafo in parrafos:
        if parrafo.strip():
            referencia_formateada = formatear_referencia(parrafo)
            referencias.append(referencia_formateada)
    referencias.sort(key=lambda ref: ref[0])
    return [ref[1] for ref in referencias]

def formatear_referencia(texto):
    match = re.match(r"([A-Za-z]+),\s([A-Za-z\.]+)\.\s\((\d{4})\)\.\s(.+)", texto)
    if match:
        autor, inicial, año, titulo = match.groups()
        referencia_formateada = f"{autor}, {inicial}. ({año}). *{titulo}*"
        return autor, referencia_formateada
    return "", texto

def aplicar_formato_apa_completo(doc_path, save_path, progress_callback):
    try:
        doc = Document(doc_path)

        # Configuración de márgenes
        for section in doc.sections:
            section.top_margin = Pt(72)
            section.bottom_margin = Pt(72)
            section.left_margin = Pt(72)
            section.right_margin = Pt(72)

            # Añadir encabezado y numeración de página
            header = section.header
            if not header.is_linked_to_previous:
                header_paragraph = header.paragraphs[0]
                header_paragraph.text = "Título del documento".upper()
                header_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        en_bibliografia = False
        referencias = []

        for i, paragraph in enumerate(doc.paragraphs):
            progress_callback(i / len(doc.paragraphs) * 100)

            # Detectar sección de bibliografía
            if paragraph.text.strip().lower() in ["referencias", "bibliografía"]:
                en_bibliografia = True
                paragraph.alignment = 0
                continue

            # Configuración de fuente y espaciado
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                r = run._element
                r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

            paragraph.paragraph_format.line_spacing = 2.0

            # Configuración de títulos principales (negritas y centrados)
            es_titulo = any(run.font.bold for run in paragraph.runs)
            if es_titulo:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Configurar referencias
            if en_bibliografia:
                referencias.append(paragraph.text)
                paragraph.clear()

        if referencias:
            referencias_ordenadas = formatear_y_organizar_referencias(referencias)
            for referencia in referencias_ordenadas:
                ref_paragraph = doc.add_paragraph(referencia)
                ref_paragraph.paragraph_format.line_spacing = 2.0
                ref_paragraph.paragraph_format.first_line_indent = None
                ref_paragraph.paragraph_format.left_indent = Cm(1.27)

        doc.save(save_path)
        messagebox.showinfo("Éxito", f"Documento guardado en formato APA:\n{save_path}")
    except Exception as e:
        messagebox.showerror("Error", f"Hubo un error: {e}")

def cargar_documento(progress_bar):
    doc_path = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")])
    if doc_path:
        save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Documents", "*.docx")])
        if save_path:
            progress_bar['value'] = 0
            aplicar_formato_apa_completo(doc_path, save_path, lambda progress: progress_bar.step(progress))
            progress_bar['value'] = 100

def iniciar_interfaz():
    root = tk.Tk()
    root.title("Formato APA Completo")
    root.geometry("500x300")
    root.configure(bg="#f0f0f0")

    label = tk.Label(root, text="Formato APA 7ª Edición - Completo", font=("Helvetica", 16), bg="#f0f0f0")
    label.pack(pady=20)

    frame = tk.Frame(root, bg="#f0f0f0")
    frame.pack(pady=10)

    progress_bar = ttk.Progressbar(root, orient="horizontal", length=400, mode="determinate")
    progress_bar.pack(pady=20)

    boton_cargar = ttk.Button(frame, text="Abrir Documento", command=lambda: cargar_documento(progress_bar))
    boton_cargar.grid(row=0, column=0, padx=10)

    boton_salir = ttk.Button(frame, text="Salir", command=root.destroy)
    boton_salir.grid(row=0, column=1, padx=10)

    root.mainloop()

if __name__ == "__main__":
    iniciar_interfaz()
